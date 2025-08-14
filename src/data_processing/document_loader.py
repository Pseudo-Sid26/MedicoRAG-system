# import os
# import PyPDF2
# import docx
# import pandas as pd
# from typing import List, Dict, Any
# import streamlit as st
# from pathlib import Path
# import logging
#
# logger = logging.getLogger(__name__)
#
#
# class DocumentLoader:
#     """Handles loading of various document types for medical literature"""
#
#     def __init__(self):
#         self.supported_formats = ['.pdf', '.txt', '.docx', '.csv']
#
#     def load_pdf(self, file_path: str) -> str:
#         """Load and extract text from PDF files"""
#         try:
#             text = ""
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages:
#                     text += page.extract_text() + "\n"
#             return text
#         except Exception as e:
#             logger.error(f"Error loading PDF {file_path}: {str(e)}")
#             raise
#
#     def load_docx(self, file_path: str) -> str:
#         """Load and extract text from DOCX files"""
#         try:
#             doc = docx.Document(file_path)
#             text = ""
#             for paragraph in doc.paragraphs:
#                 text += paragraph.text + "\n"
#             return text
#         except Exception as e:
#             logger.error(f"Error loading DOCX {file_path}: {str(e)}")
#             raise
#
#     def load_txt(self, file_path: str) -> str:
#         """Load text from TXT files"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 return file.read()
#         except Exception as e:
#             logger.error(f"Error loading TXT {file_path}: {str(e)}")
#             raise
#
#     def load_csv(self, file_path: str) -> str:
#         """Load and convert CSV to text format"""
#         try:
#             df = pd.read_csv(file_path)
#             # Convert CSV to structured text
#             text = f"Dataset: {os.path.basename(file_path)}\n"
#             text += f"Columns: {', '.join(df.columns)}\n"
#             text += f"Rows: {len(df)}\n\n"
#
#             # Add sample data
#             text += "Sample Data:\n"
#             text += df.head(10).to_string(index=False)
#
#             return text
#         except Exception as e:
#             logger.error(f"Error loading CSV {file_path}: {str(e)}")
#             raise
#
#     def load_uploaded_file(self, uploaded_file) -> Dict[str, Any]:
#         """Load file uploaded through Streamlit"""
#         try:
#             file_extension = Path(uploaded_file.name).suffix.lower()
#
#             if file_extension == '.pdf':
#                 # Save uploaded file temporarily
#                 with open(f"temp_{uploaded_file.name}", "wb") as f:
#                     f.write(uploaded_file.getvalue())
#                 text = self.load_pdf(f"temp_{uploaded_file.name}")
#                 os.remove(f"temp_{uploaded_file.name}")
#
#             elif file_extension == '.docx':
#                 with open(f"temp_{uploaded_file.name}", "wb") as f:
#                     f.write(uploaded_file.getvalue())
#                 text = self.load_docx(f"temp_{uploaded_file.name}")
#                 os.remove(f"temp_{uploaded_file.name}")
#
#             elif file_extension == '.txt':
#                 text = uploaded_file.getvalue().decode('utf-8')
#
#             elif file_extension == '.csv':
#                 df = pd.read_csv(uploaded_file)
#                 text = f"Dataset: {uploaded_file.name}\n"
#                 text += f"Columns: {', '.join(df.columns)}\n"
#                 text += f"Rows: {len(df)}\n\n"
#                 text += "Sample Data:\n"
#                 text += df.head(10).to_string(index=False)
#
#             else:
#                 raise ValueError(f"Unsupported file format: {file_extension}")
#
#             return {
#                 'filename': uploaded_file.name,
#                 'content': text,
#                 'size': len(text),
#                 'type': file_extension[1:]  # Remove the dot
#             }
#
#         except Exception as e:
#             logger.error(f"Error loading uploaded file {uploaded_file.name}: {str(e)}")
#             raise
#
#     def load_directory(self, directory_path: str) -> List[Dict[str, Any]]:
#         """Load all supported documents from a directory"""
#         documents = []
#         directory = Path(directory_path)
#
#         for file_path in directory.rglob('*'):
#             if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
#                 try:
#                     if file_path.suffix.lower() == '.pdf':
#                         content = self.load_pdf(str(file_path))
#                     elif file_path.suffix.lower() == '.docx':
#                         content = self.load_docx(str(file_path))
#                     elif file_path.suffix.lower() == '.txt':
#                         content = self.load_txt(str(file_path))
#                     elif file_path.suffix.lower() == '.csv':
#                         content = self.load_csv(str(file_path))
#
#                     documents.append({
#                         'filename': file_path.name,
#                         'filepath': str(file_path),
#                         'content': content,
#                         'size': len(content),
#                         'type': file_path.suffix[1:]  # Remove the dot
#                     })
#
#                 except Exception as e:
#                     logger.warning(f"Skipped {file_path} due to error: {str(e)}")
#                     continue
#
#         return documents
#
#     def validate_medical_document(self, content: str) -> Dict[str, Any]:
#         """Validate if document contains medical content"""
#         medical_keywords = [
#             'patient', 'diagnosis', 'treatment', 'symptom', 'medication',
#             'clinical', 'medical', 'therapy', 'disease', 'condition',
#             'healthcare', 'physician', 'doctor', 'nurse', 'hospital',
#             'study', 'research', 'trial', 'evidence', 'guidelines'
#         ]
#
#         content_lower = content.lower()
#         found_keywords = [kw for kw in medical_keywords if kw in content_lower]
#
#         return {
#             'is_medical': len(found_keywords) >= 3,
#             'medical_keywords_found': found_keywords,
#             'confidence_score': len(found_keywords) / len(medical_keywords)
#         }
#
#     def validate_document_type(self, content: str, filename: str) -> str:
#         """Determine if document is journal, guideline, or patient data"""
#         content_lower = content.lower()
#         filename_lower = filename.lower()
#
#         # Clinical guidelines indicators
#         if any(term in content_lower for term in ['guideline', 'recommendation', 'protocol', 'consensus']):
#             return 'clinical_guideline'
#
#         # Medical journal indicators
#         elif any(term in content_lower for term in ['abstract', 'methods', 'results', 'discussion', 'doi']):
#             return 'medical_journal'
#
#         # Patient data indicators
#         elif any(term in content_lower for term in ['patient', 'mrn', 'admission', 'discharge']):
#             return 'patient_data'
#
#         return 'unknown'
#
#     def assess_regulatory_compliance(self, content: str) -> Dict[str, Any]:
#         """Check regulatory approval status"""
#         content_lower = content.lower()
#
#         compliance_status = {
#             'fda_approved': 'fda approved' in content_lower or 'fda clearance' in content_lower,
#             'ema_approved': 'ema approved' in content_lower,
#             'investigational': 'investigational' in content_lower or 'clinical trial' in content_lower,
#             'off_label': 'off-label' in content_lower,
#             'regulatory_warnings': []
#         }
#
#         # Check for regulatory warnings
#         warning_terms = ['black box warning', 'contraindicated', 'not approved']
#         for term in warning_terms:
#             if term in content_lower:
#                 compliance_status['regulatory_warnings'].append(term)
#
#         return compliance_status

import os
import re
import tempfile
import time
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, BinaryIO
import logging
from functools import lru_cache
import concurrent.futures
from dataclasses import dataclass
import hashlib

# Enhanced imports with fallbacks
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    try:
        import pypdf as PyPDF2

        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd

    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

import streamlit as st

logger = logging.getLogger(__name__)


@dataclass
class DocumentMetrics:
    """Data class for document processing metrics"""
    files_processed: int
    total_size: int
    processing_time: float
    errors: int
    medical_documents: int
    avg_confidence: float


@dataclass
class DocumentInfo:
    """Enhanced document information structure"""
    filename: str
    content: str
    size: int
    file_type: str
    document_type: str
    medical_confidence: float
    processing_time: float
    metadata: Dict[str, Any]
    quality_score: float


class DocumentLoader:
    """Optimized document loader with enhanced medical document processing"""

    def __init__(self, enable_caching: bool = True, max_file_size: int = 100 * 1024 * 1024):
        self.supported_formats = self._get_supported_formats()
        self.enable_caching = enable_caching
        self.max_file_size = max_file_size

        # Performance: Cache for processed files
        self._content_cache = {} if enable_caching else None
        self._validation_cache = {} if enable_caching else None

        # Compile regex patterns for better performance
        self._compile_patterns()

        # Track processing metrics
        self.metrics = {
            'files_processed': 0,
            'total_processing_time': 0,
            'cache_hits': 0,
            'errors': 0
        }

    def _get_supported_formats(self) -> List[str]:
        """Get supported formats based on available libraries"""
        formats = ['.txt']

        if PDF_AVAILABLE:
            formats.append('.pdf')
        if DOCX_AVAILABLE:
            formats.append('.docx')
        if PANDAS_AVAILABLE:
            formats.extend(['.csv', '.xlsx'])

        return formats

    def _compile_patterns(self):
        """Compile regex patterns for better performance"""
        # Medical content patterns
        self.medical_patterns = {
            'keywords': re.compile(
                r'\b(?:patient|diagnosis|treatment|symptom|medication|clinical|medical|therapy|'
                r'disease|condition|healthcare|physician|doctor|nurse|hospital|study|research|'
                r'trial|evidence|guidelines|therapeutic|pharmaceutical|pathology|radiology|'
                r'cardiology|neurology|oncology|surgery|anesthesia|emergency)\b',
                re.IGNORECASE
            ),
            'drug_names': re.compile(
                r'\b\w+(?:cillin|mycin|azole|pril|sartan|statin|ide|ine|ol|al|an|er|in|um|ate|ase)\b',
                re.IGNORECASE
            ),
            'measurements': re.compile(
                r'\b\d+(?:\.\d+)?\s*(?:mg|ml|mmHg|°[CF]|bpm|kg|cm|mm|hrs?|min|sec|units?|iu|mcg)\b',
                re.IGNORECASE
            ),
            'medical_codes': re.compile(
                r'\b(?:ICD-?10?|CPT|SNOMED|LOINC|NDC)[-\s]?\w+\b',
                re.IGNORECASE
            )
        }

        # Document type patterns
        self.document_type_patterns = {
            'clinical_guideline': re.compile(
                r'\b(?:guideline|recommendation|protocol|consensus|standard|policy|'
                r'best practice|clinical pathway)\b',
                re.IGNORECASE
            ),
            'medical_journal': re.compile(
                r'\b(?:abstract|methods?|results?|discussion|conclusion|doi|pmid|'
                r'introduction|background|objective|study design|clinical trial)\b',
                re.IGNORECASE
            ),
            'patient_data': re.compile(
                r'\b(?:patient|mrn|medical record|admission|discharge|hospital stay|'
                r'chief complaint|history of present illness|physical exam)\b',
                re.IGNORECASE
            ),
            'drug_information': re.compile(
                r'\b(?:dosage|administration|contraindications|side effects|interactions|'
                r'pharmacokinetics|pharmacodynamics|mechanism of action)\b',
                re.IGNORECASE
            ),
            'regulatory': re.compile(
                r'\b(?:fda approved?|ema approved?|regulatory|compliance|label|'
                r'black box warning|contraindicated|clinical trial phase)\b',
                re.IGNORECASE
            )
        }

    def _get_file_hash(self, file_content: bytes) -> str:
        """Generate hash for file content"""
        return hashlib.md5(file_content).hexdigest()

    def _validate_file_size(self, file_size: int, filename: str) -> bool:
        """Validate file size"""
        if file_size > self.max_file_size:
            logger.warning(f"File {filename} exceeds maximum size ({file_size} > {self.max_file_size})")
            return False
        return True

    @lru_cache(maxsize=100)
    def _cached_text_extraction(self, file_hash: str, file_type: str, content_preview: str) -> str:
        """Cache text extraction results"""
        # This is a placeholder for actual extraction logic
        # The real extraction happens in the main methods
        return content_preview

    def load_pdf(self, file_path: str) -> str:
        """Optimized PDF loading with enhanced error handling"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 or pypdf not available. Install with: pip install PyPDF2")

        try:
            start_time = time.time()
            text = ""

            with open(file_path, 'rb') as file:
                # Check file size first
                file.seek(0, 2)  # Seek to end
                file_size = file.tell()
                file.seek(0)  # Seek back to beginning

                if not self._validate_file_size(file_size, file_path):
                    raise ValueError(f"File too large: {file_size} bytes")

                # Check if cached
                if self.enable_caching:
                    file_hash = self._get_file_hash(file.read())
                    file.seek(0)

                    if file_hash in self._content_cache:
                        self.metrics['cache_hits'] += 1
                        return self._content_cache[file_hash]

                # Extract text
                try:
                    pdf_reader = PyPDF2.PdfReader(file)

                    # Check if PDF is encrypted
                    if pdf_reader.is_encrypted:
                        logger.warning(f"PDF {file_path} is encrypted, attempting to decrypt")
                        try:
                            pdf_reader.decrypt("")  # Try empty password
                        except Exception:
                            raise ValueError("PDF is password protected")

                    # Extract text from pages with progress tracking
                    total_pages = len(pdf_reader.pages)
                    for i, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                        except Exception as e:
                            logger.warning(f"Error extracting page {i + 1} from {file_path}: {str(e)}")
                            continue

                    # Fallback for PDFs with poor text extraction
                    if not text.strip():
                        logger.warning(f"No text extracted from {file_path}, might be image-based PDF")
                        text = f"[PDF file: {os.path.basename(file_path)} - {total_pages} pages - Text extraction failed]"

                except Exception as e:
                    # Try alternative extraction method
                    try:
                        file.seek(0)
                        pdf_reader = PyPDF2.PdfReader(file, strict=False)
                        for page in pdf_reader.pages:
                            try:
                                text += page.extract_text() + "\n"
                            except:
                                continue
                    except Exception:
                        raise ValueError(f"Unable to process PDF: {str(e)}")

                # Cache result
                if self.enable_caching and file_hash:
                    self._content_cache[file_hash] = text

                processing_time = time.time() - start_time
                logger.debug(f"PDF processing completed in {processing_time:.2f}s")

                return text

        except Exception as e:
            self.metrics['errors'] += 1
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise

    def load_docx(self, file_path: str) -> str:
        """Optimized DOCX loading with enhanced content extraction"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")

        try:
            start_time = time.time()

            # Check file size
            file_size = os.path.getsize(file_path)
            if not self._validate_file_size(file_size, file_path):
                raise ValueError(f"File too large: {file_size} bytes")

            # Check cache
            if self.enable_caching:
                with open(file_path, 'rb') as f:
                    file_hash = self._get_file_hash(f.read())

                if file_hash in self._content_cache:
                    self.metrics['cache_hits'] += 1
                    return self._content_cache[file_hash]

            doc = docx.Document(file_path)
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))

                if table_text:
                    text_parts.append("\n".join(table_text))

            # Extract headers and footers
            for section in doc.sections:
                if section.header:
                    header_text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
                    if header_text:
                        text_parts.insert(0, f"Header: {header_text}")

                if section.footer:
                    footer_text = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip())
                    if footer_text:
                        text_parts.append(f"Footer: {footer_text}")

            text = "\n\n".join(text_parts)

            # Cache result
            if self.enable_caching and file_hash:
                self._content_cache[file_hash] = text

            processing_time = time.time() - start_time
            logger.debug(f"DOCX processing completed in {processing_time:.2f}s")

            return text

        except Exception as e:
            self.metrics['errors'] += 1
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise

    def load_txt(self, file_path: str, encoding: str = 'utf-8') -> str:
        """Enhanced text file loading with encoding detection"""
        try:
            start_time = time.time()

            # Check file size
            file_size = os.path.getsize(file_path)
            if not self._validate_file_size(file_size, file_path):
                raise ValueError(f"File too large: {file_size} bytes")

            # Try multiple encodings
            encodings = [encoding, 'utf-8', 'latin1', 'cp1252', 'ascii']

            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as file:
                        content = file.read()

                    processing_time = time.time() - start_time
                    logger.debug(f"TXT processing completed in {processing_time:.2f}s with encoding {enc}")
                    return content

                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logger.warning(f"Error with encoding {enc}: {str(e)}")
                    continue

            raise ValueError("Unable to decode file with any supported encoding")

        except Exception as e:
            self.metrics['errors'] += 1
            logger.error(f"Error loading TXT {file_path}: {str(e)}")
            raise

    def load_csv(self, file_path: str) -> str:
        """Enhanced CSV loading with better formatting and error handling"""
        if not PANDAS_AVAILABLE:
            raise ImportError("pandas not available. Install with: pip install pandas")

        try:
            start_time = time.time()

            # Check file size
            file_size = os.path.getsize(file_path)
            if not self._validate_file_size(file_size, file_path):
                raise ValueError(f"File too large: {file_size} bytes")

            # Try different separators and encodings
            separators = [',', ';', '\t', '|']
            encodings = ['utf-8', 'latin1', 'cp1252']

            df = None
            for sep in separators:
                for enc in encodings:
                    try:
                        df = pd.read_csv(file_path, sep=sep, encoding=enc, nrows=1000)  # Limit rows for performance
                        if len(df.columns) > 1:  # Successfully parsed with multiple columns
                            break
                    except Exception:
                        continue
                if df is not None and len(df.columns) > 1:
                    break

            if df is None:
                # Fallback: read as plain text
                return self.load_txt(file_path)

            # Generate structured text representation
            text_parts = [
                f"Dataset: {os.path.basename(file_path)}",
                f"Columns ({len(df.columns)}): {', '.join(str(col) for col in df.columns)}",
                f"Rows: {len(df)}",
                ""
            ]

            # Add data type information
            if not df.empty:
                text_parts.append("Data Types:")
                for col in df.columns:
                    dtype = str(df[col].dtype)
                    non_null = df[col].count()
                    text_parts.append(f"  {col}: {dtype} ({non_null} non-null)")
                text_parts.append("")

            # Add sample data with better formatting
            if not df.empty:
                text_parts.append("Sample Data:")
                sample_size = min(10, len(df))

                # Format data for readability
                for i in range(sample_size):
                    row_text = f"Row {i + 1}:"
                    for col in df.columns:
                        value = df.iloc[i][col]
                        if pd.isna(value):
                            value = "[NULL]"
                        row_text += f"\n  {col}: {value}"
                    text_parts.append(row_text)
                    text_parts.append("")

            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_parts.append("Numeric Summary:")
                for col in numeric_cols:
                    stats = df[col].describe()
                    text_parts.append(
                        f"  {col}: mean={stats['mean']:.2f}, std={stats['std']:.2f}, min={stats['min']}, max={stats['max']}")
                text_parts.append("")

            text = "\n".join(text_parts)

            processing_time = time.time() - start_time
            logger.debug(f"CSV processing completed in {processing_time:.2f}s")

            return text

        except Exception as e:
            self.metrics['errors'] += 1
            logger.error(f"Error loading CSV {file_path}: {str(e)}")
            raise

    def load_uploaded_file(self, uploaded_file) -> Dict[str, Any]:
        """Optimized uploaded file processing with enhanced metadata extraction"""
        try:
            start_time = time.time()
            file_extension = Path(uploaded_file.name).suffix.lower()

            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")

            # Check file size
            file_content = uploaded_file.getvalue()
            if not self._validate_file_size(len(file_content), uploaded_file.name):
                raise ValueError(f"File too large: {len(file_content)} bytes")

            # Generate file hash for caching
            file_hash = self._get_file_hash(file_content) if self.enable_caching else None

            # Check cache
            if self.enable_caching and file_hash in self._content_cache:
                self.metrics['cache_hits'] += 1
                cached_result = self._content_cache[file_hash].copy()
                cached_result['filename'] = uploaded_file.name  # Update filename
                return cached_result

            # Process file based on type
            text = ""
            metadata = {}

            if file_extension == '.pdf':
                text = self._process_pdf_upload(uploaded_file, file_content)
            elif file_extension == '.docx':
                text = self._process_docx_upload(uploaded_file, file_content)
            elif file_extension == '.txt':
                text = self._process_txt_upload(file_content)
            elif file_extension in ['.csv', '.xlsx']:
                text, metadata = self._process_spreadsheet_upload(uploaded_file, file_content)

            # Validate and enhance content
            medical_validation = self.validate_medical_document(text)
            document_type = self.validate_document_type(text, uploaded_file.name)
            quality_score = self._assess_document_quality(text)

            result = DocumentInfo(
                filename=uploaded_file.name,
                content=text,
                size=len(text),
                file_type=file_extension[1:],
                document_type=document_type,
                medical_confidence=medical_validation['confidence_score'],
                processing_time=time.time() - start_time,
                metadata={
                    **metadata,
                    'original_file_size': len(file_content),
                    'medical_keywords': medical_validation['medical_keywords_found'],
                    'regulatory_info': self.assess_regulatory_compliance(text)
                },
                quality_score=quality_score
            ).__dict__

            # Cache result
            if self.enable_caching and file_hash:
                self._content_cache[file_hash] = result.copy()

            self.metrics['files_processed'] += 1
            self.metrics['total_processing_time'] += result['processing_time']

            return result

        except Exception as e:
            self.metrics['errors'] += 1
            logger.error(f"Error loading uploaded file {uploaded_file.name}: {str(e)}")
            raise

    def _process_pdf_upload(self, uploaded_file, file_content: bytes) -> str:
        """Process uploaded PDF file"""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file.flush()

            try:
                text = self.load_pdf(temp_file.name)
            finally:
                os.unlink(temp_file.name)

            return text

    def _process_docx_upload(self, uploaded_file, file_content: bytes) -> str:
        """Process uploaded DOCX file"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file.flush()

            try:
                text = self.load_docx(temp_file.name)
            finally:
                os.unlink(temp_file.name)

            return text

    def _process_txt_upload(self, file_content: bytes) -> str:
        """Process uploaded text file"""
        # Try different encodings
        encodings = ['utf-8', 'latin1', 'cp1252', 'ascii']

        for encoding in encodings:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue

        raise ValueError("Unable to decode text file with any supported encoding")

    def _process_spreadsheet_upload(self, uploaded_file, file_content: bytes) -> tuple[str, Dict[str, Any]]:
        """Process uploaded CSV/Excel file"""
        if not PANDAS_AVAILABLE:
            raise ImportError("pandas not available for spreadsheet processing")

        metadata = {}

        try:
            if uploaded_file.name.endswith('.csv'):
                # Reset file pointer for pandas
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file, nrows=1000)  # Limit for performance
            else:  # Excel file
                with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as temp_file:
                    temp_file.write(file_content)
                    temp_file.flush()

                    try:
                        df = pd.read_excel(temp_file.name, nrows=1000)
                    finally:
                        os.unlink(temp_file.name)

            # Extract metadata
            metadata = {
                'sheet_name': getattr(df, 'name', 'Sheet1'),
                'column_count': len(df.columns),
                'row_count': len(df),
                'numeric_columns': list(df.select_dtypes(include=['number']).columns),
                'text_columns': list(df.select_dtypes(include=['object']).columns),
                'missing_data_summary': df.isnull().sum().to_dict()
            }

            # Generate text representation (reuse CSV logic)
            with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False) as temp_csv:
                df.to_csv(temp_csv.name, index=False)
                try:
                    text = self.load_csv(temp_csv.name)
                finally:
                    os.unlink(temp_csv.name)

            return text, metadata

        except Exception as e:
            logger.error(f"Error processing spreadsheet: {str(e)}")
            # Fallback to text processing
            text = self._process_txt_upload(file_content)
            return text, metadata

    def load_directory(self, directory_path: str, parallel: bool = True) -> List[Dict[str, Any]]:
        """Optimized directory loading with parallel processing"""
        directory = Path(directory_path)
        if not directory.exists():
            raise ValueError(f"Directory does not exist: {directory_path}")

        # Find all supported files
        file_paths = []
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                file_paths.append(file_path)

        if not file_paths:
            logger.warning(f"No supported files found in {directory_path}")
            return []

        documents = []
        start_time = time.time()

        if parallel and len(file_paths) > 1:
            # Parallel processing for multiple files
            with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                future_to_path = {
                    executor.submit(self._load_single_file, file_path): file_path
                    for file_path in file_paths
                }

                for future in concurrent.futures.as_completed(future_to_path):
                    file_path = future_to_path[future]
                    try:
                        doc_info = future.result()
                        if doc_info:
                            documents.append(doc_info)
                    except Exception as e:
                        logger.warning(f"Skipped {file_path} due to error: {str(e)}")
                        continue
        else:
            # Sequential processing
            for file_path in file_paths:
                try:
                    doc_info = self._load_single_file(file_path)
                    if doc_info:
                        documents.append(doc_info)
                except Exception as e:
                    logger.warning(f"Skipped {file_path} due to error: {str(e)}")
                    continue

        total_time = time.time() - start_time
        logger.info(f"Loaded {len(documents)} documents from {directory_path} in {total_time:.2f}s")

        return documents

    def _load_single_file(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """Load a single file with error handling"""
        try:
            start_time = time.time()

            if file_path.suffix.lower() == '.pdf':
                content = self.load_pdf(str(file_path))
            elif file_path.suffix.lower() == '.docx':
                content = self.load_docx(str(file_path))
            elif file_path.suffix.lower() == '.txt':
                content = self.load_txt(str(file_path))
            elif file_path.suffix.lower() in ['.csv', '.xlsx']:
                content = self.load_csv(str(file_path))
            else:
                return None

            # Enhanced document analysis
            medical_validation = self.validate_medical_document(content)
            document_type = self.validate_document_type(content, file_path.name)
            quality_score = self._assess_document_quality(content)

            return {
                'filename': file_path.name,
                'filepath': str(file_path),
                'content': content,
                'size': len(content),
                'type': file_path.suffix[1:],
                'document_type': document_type,
                'medical_confidence': medical_validation['confidence_score'],
                'processing_time': time.time() - start_time,
                'quality_score': quality_score,
                'metadata': {
                    'medical_keywords': medical_validation['medical_keywords_found'],
                    'regulatory_info': self.assess_regulatory_compliance(content),
                    'file_stats': {
                        'word_count': len(content.split()),
                        'char_count': len(content),
                        'line_count': len(content.split('\n'))
                    }
                }
            }

        except Exception as e:
            logger.error(f"Error loading file {file_path}: {str(e)}")
            return None

    def validate_medical_document(self, content: str) -> Dict[str, Any]:
        """Enhanced medical document validation with compiled patterns"""
        if self.enable_caching:
            content_hash = str(hash(content))
            if content_hash in self._validation_cache:
                return self._validation_cache[content_hash]

        content_lower = content.lower()

        # Count different types of medical content
        medical_scores = {}
        total_matches = 0

        for pattern_name, pattern in self.medical_patterns.items():
            matches = pattern.findall(content)
            medical_scores[pattern_name] = len(matches)
            total_matches += len(matches)

        # Calculate confidence based on content length and match density
        content_length = len(content.split())
        density = total_matches / max(content_length, 1)

        # Base confidence from keyword presence
        base_confidence = min(total_matches / 50.0, 1.0)  # Normalize to 50 expected matches

        # Boost confidence based on density
        density_boost = min(density * 10, 0.3)  # Max 30% boost

        # Boost for specific medical indicators
        specialty_boost = 0
        specialty_terms = ['clinical trial', 'evidence-based', 'peer-reviewed', 'systematic review']
        for term in specialty_terms:
            if term in content_lower:
                specialty_boost += 0.05

        final_confidence = min(base_confidence + density_boost + specialty_boost, 1.0)

        # Extract found keywords for reporting
        found_keywords = []
        for pattern_name, pattern in self.medical_patterns.items():
            matches = pattern.findall(content)
            found_keywords.extend(matches[:5])  # Limit to 5 per pattern

        result = {
            'is_medical': final_confidence >= 0.3,
            'medical_keywords_found': list(set(found_keywords))[:20],  # Limit total keywords
            'confidence_score': final_confidence,
            'medical_scores': medical_scores,
            'content_density': density
        }

        # Cache result
        if self.enable_caching:
            self._validation_cache[content_hash] = result

        return result

    def validate_document_type(self, content: str, filename: str) -> str:
        """Enhanced document type detection"""
        content_lower = content.lower()
        filename_lower = filename.lower()

        # Score each document type
        type_scores = {}

        for doc_type, pattern in self.document_type_patterns.items():
            matches = pattern.findall(content_lower)
            # Weight by match count and normalize by content length
            score = len(matches) / max(len(content.split()) / 1000, 1)
            type_scores[doc_type] = score

        # Additional filename-based hints
        filename_hints = {
            'clinical_guideline': ['guideline', 'protocol', 'standard', 'recommendation'],
            'medical_journal': ['journal', 'article', 'paper', 'study'],
            'patient_data': ['patient', 'record', 'chart', 'case'],
            'drug_information': ['drug', 'medication', 'pharma', 'label'],
            'regulatory': ['fda', 'ema', 'regulatory', 'approval']
        }

        for doc_type, hints in filename_hints.items():
            for hint in hints:
                if hint in filename_lower:
                    type_scores[doc_type] = type_scores.get(doc_type, 0) + 0.1

        # Find the document type with highest score
        if type_scores:
            best_type = max(type_scores.items(), key=lambda x: x[1])
            if best_type[1] > 0.1:  # Minimum threshold
                return best_type[0]

        return 'unknown'

    def _assess_document_quality(self, content: str) -> float:
        """Assess document quality based on various factors"""
        if not content.strip():
            return 0.0

        quality_score = 0.5  # Base score

        # Content length factor
        word_count = len(content.split())
        if 100 <= word_count <= 10000:
            quality_score += 0.2
        elif word_count < 50:
            quality_score -= 0.2

        # Structure indicators
        if re.search(r'\b(abstract|introduction|conclusion)\b', content, re.IGNORECASE):
            quality_score += 0.1

        # Reference indicators
        if re.search(r'\[\d+\]|\(\d{4}\)|doi:', content, re.IGNORECASE):
            quality_score += 0.1

        # Medical terminology density
        medical_terms = self.medical_patterns['keywords'].findall(content)
        term_density = len(medical_terms) / max(word_count, 1)
        if 0.02 <= term_density <= 0.15:  # Optimal range
            quality_score += 0.1

        # Readability indicators (sentences, paragraphs)
        sentences = len(re.split(r'[.!?]+', content))
        paragraphs = len(content.split('\n\n'))
        if sentences > 5 and paragraphs > 1:
            quality_score += 0.1

        return min(quality_score, 1.0)

    def assess_regulatory_compliance(self, content: str) -> Dict[str, Any]:
        """Enhanced regulatory compliance assessment"""
        content_lower = content.lower()

        compliance_status = {
            'fda_approved': bool(re.search(r'\bfda\s+approved?\b', content_lower)),
            'ema_approved': bool(re.search(r'\bema\s+approved?\b', content_lower)),
            'health_canada_approved': bool(re.search(r'\bhealth\s+canada\s+approved?\b', content_lower)),
            'investigational': bool(
                re.search(r'\b(?:investigational|clinical\s+trial|phase\s+[i-iv])\b', content_lower)),
            'off_label': bool(re.search(r'\boff[-\s]label\b', content_lower)),
            'orphan_drug': bool(re.search(r'\borphan\s+drug\b', content_lower)),
            'breakthrough_therapy': bool(re.search(r'\bbreakthrough\s+therapy\b', content_lower)),
            'fast_track': bool(re.search(r'\bfast\s+track\b', content_lower)),
            'regulatory_warnings': [],
            'approval_dates': [],
            'regulatory_agencies': []
        }

        # Check for regulatory warnings
        warning_patterns = {
            'black_box_warning': r'\bblack\s+box\s+warning\b',
            'contraindicated': r'\bcontraindicated?\b',
            'not_approved': r'\bnot\s+approved?\b',
            'recalled': r'\brecalled?\b',
            'withdrawn': r'\bwithdrawn\b',
            'safety_alert': r'\bsafety\s+alert\b'
        }

        for warning_type, pattern in warning_patterns.items():
            if re.search(pattern, content_lower):
                compliance_status['regulatory_warnings'].append(warning_type)

        # Extract approval dates
        date_patterns = [
            r'\bapproved?\s+(?:on\s+|in\s+)?(\d{1,2}[/-]\d{1,2}[/-]\d{4})\b',
            r'\bapproved?\s+(?:on\s+|in\s+)?(\w+\s+\d{1,2},?\s+\d{4})\b',
            r'\bapproved?\s+(?:on\s+|in\s+)?(\w+\s+\d{4})\b'
        ]

        for pattern in date_patterns:
            matches = re.findall(pattern, content_lower)
            compliance_status['approval_dates'].extend(matches)

        # Extract regulatory agencies mentioned
        agency_patterns = {
            'FDA': r'\bfda\b',
            'EMA': r'\bema\b',
            'Health Canada': r'\bhealth\s+canada\b',
            'PMDA': r'\bpmda\b',
            'TGA': r'\btga\b',
            'ANVISA': r'\banvisa\b',
            'NMPA': r'\bnmpa\b'
        }

        for agency, pattern in agency_patterns.items():
            if re.search(pattern, content_lower):
                compliance_status['regulatory_agencies'].append(agency)

        return compliance_status

    def get_processing_metrics(self) -> DocumentMetrics:
        """Get document processing performance metrics"""
        avg_time = (self.metrics['total_processing_time'] /
                    max(self.metrics['files_processed'], 1))

        return DocumentMetrics(
            files_processed=self.metrics['files_processed'],
            total_size=0,  # Would need to track this separately
            processing_time=self.metrics['total_processing_time'],
            errors=self.metrics['errors'],
            medical_documents=0,  # Would need to track this separately
            avg_confidence=0.0  # Would need to track this separately
        )

    def clear_cache(self):
        """Clear all caches to free memory"""
        if self._content_cache:
            self._content_cache.clear()
        if self._validation_cache:
            self._validation_cache.clear()

        # Clear LRU cache
        self._cached_text_extraction.cache_clear()

        logger.info("Document loader caches cleared")

    def get_cache_stats(self) -> Dict[str, Any]:
        """Get cache statistics"""
        return {
            'content_cache_size': len(self._content_cache) if self._content_cache else 0,
            'validation_cache_size': len(self._validation_cache) if self._validation_cache else 0,
            'cache_hits': self.metrics['cache_hits'],
            'lru_cache_info': self._cached_text_extraction.cache_info()._asdict()
        }

    def batch_process_files(self, file_paths: List[str], batch_size: int = 5) -> List[Dict[str, Any]]:
        """Process multiple files in batches for better memory management"""
        results = []
        total_files = len(file_paths)

        for i in range(0, total_files, batch_size):
            batch = file_paths[i:i + batch_size]
            batch_start_time = time.time()

            # Process batch in parallel
            with concurrent.futures.ThreadPoolExecutor(max_workers=min(batch_size, 3)) as executor:
                future_to_path = {
                    executor.submit(self._load_single_file, Path(file_path)): file_path
                    for file_path in batch
                }

                batch_results = []
                for future in concurrent.futures.as_completed(future_to_path):
                    file_path = future_to_path[future]
                    try:
                        result = future.result()
                        if result:
                            batch_results.append(result)
                    except Exception as e:
                        logger.error(f"Error processing {file_path}: {str(e)}")
                        continue

                results.extend(batch_results)

            batch_time = time.time() - batch_start_time
            logger.info(f"Processed batch {i // batch_size + 1}/{(total_files - 1) // batch_size + 1} "
                        f"({len(batch_results)}/{len(batch)} files) in {batch_time:.2f}s")

            # Optional: Clear cache between batches to manage memory
            if i > 0 and i % (batch_size * 4) == 0:  # Clear every 4 batches
                self.clear_cache()

        return results

    def validate_file_integrity(self, file_path: str) -> Dict[str, Any]:
        """Validate file integrity and provide detailed information"""
        file_path = Path(file_path)

        integrity_info = {
            'exists': file_path.exists(),
            'is_file': file_path.is_file() if file_path.exists() else False,
            'size': file_path.stat().st_size if file_path.exists() else 0,
            'extension': file_path.suffix.lower(),
            'supported': file_path.suffix.lower() in self.supported_formats,
            'readable': False,
            'corrupted': False,
            'error_message': None
        }

        if not integrity_info['exists']:
            integrity_info['error_message'] = "File does not exist"
            return integrity_info

        if not integrity_info['is_file']:
            integrity_info['error_message'] = "Path is not a file"
            return integrity_info

        if not integrity_info['supported']:
            integrity_info['error_message'] = f"Unsupported file format: {integrity_info['extension']}"
            return integrity_info

        # Test file readability
        try:
            if integrity_info['extension'] == '.pdf':
                if PDF_AVAILABLE:
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        # Try to read first page
                        if len(pdf_reader.pages) > 0:
                            pdf_reader.pages[0].extract_text()
                        integrity_info['readable'] = True
                else:
                    integrity_info['error_message'] = "PDF processing library not available"

            elif integrity_info['extension'] == '.docx':
                if DOCX_AVAILABLE:
                    doc = docx.Document(file_path)
                    # Try to access paragraphs
                    list(doc.paragraphs)
                    integrity_info['readable'] = True
                else:
                    integrity_info['error_message'] = "DOCX processing library not available"

            elif integrity_info['extension'] == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    f.read(100)  # Read first 100 characters
                integrity_info['readable'] = True

            elif integrity_info['extension'] in ['.csv', '.xlsx']:
                if PANDAS_AVAILABLE:
                    if integrity_info['extension'] == '.csv':
                        pd.read_csv(file_path, nrows=1)
                    else:
                        pd.read_excel(file_path, nrows=1)
                    integrity_info['readable'] = True
                else:
                    integrity_info['error_message'] = "Pandas library not available"

        except Exception as e:
            integrity_info['corrupted'] = True
            integrity_info['error_message'] = f"File appears corrupted: {str(e)}"

        return integrity_info

    def get_supported_formats_info(self) -> Dict[str, Any]:
        """Get information about supported formats and their capabilities"""
        return {
            'supported_formats': self.supported_formats,
            'libraries_available': {
                'pdf': PDF_AVAILABLE,
                'docx': DOCX_AVAILABLE,
                'pandas': PANDAS_AVAILABLE
            },
            'capabilities': {
                '.pdf': 'Text extraction from PDF documents' if PDF_AVAILABLE else 'Not available - install PyPDF2',
                '.docx': 'Text extraction from Word documents' if DOCX_AVAILABLE else 'Not available - install python-docx',
                '.txt': 'Plain text file reading',
                '.csv': 'Structured data processing' if PANDAS_AVAILABLE else 'Not available - install pandas',
                '.xlsx': 'Excel file processing' if PANDAS_AVAILABLE else 'Not available - install pandas'
            },
            'performance_features': [
                'Caching for repeated file processing',
                'Parallel processing for multiple files',
                'Smart encoding detection',
                'Enhanced medical content validation',
                'Regulatory compliance assessment',
                'Document quality scoring'
            ]
        }